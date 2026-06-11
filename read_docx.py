import docx
import sys

def read_docx(path):
    doc = docx.Document(path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    # Also read tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)

if __name__ == '__main__':
    path = sys.argv[1]
    text = read_docx(path)
    print(text)