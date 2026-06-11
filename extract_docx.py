import docx
import sys

def read_docx(path):
    doc = docx.Document(path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    # Also read tables
    for table_idx, table in enumerate(doc.tables):
        full_text.append(f'\n--- Table {table_idx} ---')
        for row_idx, row in enumerate(table.rows):
            row_text = []
            for cell_idx, cell in enumerate(row.cells):
                row_text.append(cell.text.strip())
            full_text.append(' | '.join(row_text))
    return '\n'.join(full_text)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <docx_file>")
        sys.exit(1)
    path = sys.argv[1]
    text = read_docx(path)
    print(text)