import docx
print("Docx imported successfully")
doc = docx.Document("統計數據-專任(案)教師數_修正.docx")
print(f"Number of paragraphs: {len(doc.paragraphs)}")
for i, para in enumerate(doc.paragraphs[:5]):  # First 5 paragraphs
    print(f"Paragraph {i}: {para.text}")