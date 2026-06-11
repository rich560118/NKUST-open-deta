import docx
import sys
import traceback

def read_docx(path):
    try:
        doc = docx.Document(path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        traceback.print_exc()
        return f"Error: {e}"

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python read_docx2.py <docx_file>")
        sys.exit(1)
    path = sys.argv[1]
    text = read_docx(path)
    print(text)